

from datetime import datetime
from logging import Logger
import os

from docx import Document
from docx.shared import Pt

from file_system_assistant.app.services.container import Container
from file_system_assistant.core.ports.file_io.file_io import FileIO
from file_system_assistant.core.ports.file_io.file_read_result import FileReadResult
from file_system_assistant.core.ports.common.models.file_metadata import FileMetadata
from file_system_assistant.core.types.result import Result


class DocxFileIO(FileIO):

    def __init__(self) -> None:
        super().__init__()
        self.logger = Container.resolve(Logger)

    def read_file(self, filepath: str) -> Result[FileReadResult, str]:
        try:
            doc = Document(filepath)

            pages: list[str] = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(pages).strip()

            stats = os.stat(filepath)

            file_metadata = FileMetadata(
                name=os.path.basename(filepath),
                size=stats.st_size,
                kind="docx",
                modified_date=datetime.fromtimestamp(stats.st_mtime),
            )

            result = FileReadResult(
                False,
                content,
                file_metadata
            )

            return Result.ok(result)

        except Exception as e:
            self.logger.error(e)
            return Result.err(f"Failed to read DOCX file")

    def write_file(self, filepath: str, content: str | bytes) -> Result[bool, str]:
        try:
            text: str = ""
            if isinstance(content, bytes):
                text = content.decode("utf-8")
            elif isinstance(content, str):
                text = content

            os.makedirs(os.path.dirname(filepath), exist_ok=True)

            if os.path.exists(filepath):
                doc = Document(filepath)
            else:
                doc = Document()

            for paragraph in text.split("\n"):
                p = doc.add_paragraph()
                run = p.add_run(paragraph)
                run.font.size = Pt(12)

            doc.save(filepath)
            return Result.ok(True)
        
        except FileNotFoundError as e:
            self.logger.error(e)
            return Result.err("File not found")
        except Exception as e:
            self.logger.error(e)
            return Result.err("Something went wrong")